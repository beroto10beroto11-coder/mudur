"""
output_generator.py — Ders programını Excel / PDF / Word olarak üretir.
"""
from __future__ import annotations

import io
from typing import Literal

from app.scheduler.data_loader import SchedulerInput
from app.scheduler.solver import SchedulerResult, ScheduledLesson

DAY_NAMES = ["Pazartesi", "Salı", "Çarşamba", "Perşembe", "Cuma"]
OutputFormat = Literal["excel", "pdf", "word"]


# ---------------------------------------------------------------------------
# Yardımcı: Hücre metni oluştur
# ---------------------------------------------------------------------------

def _cell_text(lesson: ScheduledLesson) -> str:
    return f"{lesson.course_name}\n{lesson.teacher_name}"


def _build_grid(
    entity_name: str,
    entity_type: Literal["class", "teacher"],
    result: SchedulerResult,
    daily_periods: list[int],
) -> dict[tuple[int, int], str]:
    """
    Gün × Period matrisini doldurur.
    entity_type='class'   → class_name == entity_name olanlar
    entity_type='teacher' → teacher_name == entity_name olanlar
    Döner: {(day, period): "Ders Adı\nÖğretmen Adı"}
    """
    grid: dict[tuple[int, int], str] = {}
    for lesson in result.lessons:
        if entity_type == "class" and lesson.class_name != entity_name:
            continue
        if entity_type == "teacher" and lesson.teacher_name != entity_name:
            continue
        for offset in range(lesson.length):
            period = lesson.start_period + offset
            text = _cell_text(lesson)
            grid[(lesson.day, period)] = text
    return grid


# ---------------------------------------------------------------------------
# EXCEL
# ---------------------------------------------------------------------------

def generate_excel(inp: SchedulerInput, result: SchedulerResult) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import (
        Alignment, Font, PatternFill, Border, Side, GradientFill
    )
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)  # boş default sheet sil

    # Renkler
    HEADER_FILL = PatternFill("solid", fgColor="1E3A5F")      # koyu lacivert
    DAY_FILL    = PatternFill("solid", fgColor="2563EB")      # mavi
    PERIOD_FILL = PatternFill("solid", fgColor="EFF6FF")      # açık mavi
    LESSON_FILL = PatternFill("solid", fgColor="DBEAFE")      # ders hücresi
    EMPTY_FILL  = PatternFill("solid", fgColor="F8FAFC")      # boş hücre
    WARN_FILL   = PatternFill("solid", fgColor="FEF3C7")      # uyarı

    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _write_timetable_sheet(ws, entity_name: str, entity_type: str):
        max_p = max(inp.daily_periods)
        grid = _build_grid(entity_name, entity_type, result, inp.daily_periods)

        # Sütun genişlikleri
        ws.column_dimensions["A"].width = 12
        for col in range(2, DAYS + 2):
            ws.column_dimensions[get_column_letter(col)].width = 24

        # Başlık satırı
        ws.row_dimensions[1].height = 22
        title_cell = ws.cell(row=1, column=1, value=entity_name)
        title_cell.fill = HEADER_FILL
        title_cell.font = Font(bold=True, color="FFFFFF", size=12)
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=DAYS + 1)

        # Gün başlıkları
        ws.row_dimensions[2].height = 20
        ws.cell(row=2, column=1, value="Ders / Gün").fill = PERIOD_FILL
        ws.cell(row=2, column=1).font = Font(bold=True, size=10)
        ws.cell(row=2, column=1).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row=2, column=1).border = border
        for d, day_name in enumerate(DAY_NAMES):
            cell = ws.cell(row=2, column=d + 2, value=day_name)
            cell.fill = DAY_FILL
            cell.font = Font(bold=True, color="FFFFFF", size=10)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border

        # Period satırları
        for p in range(1, max_p + 1):
            row = p + 2
            ws.row_dimensions[row].height = 50

            period_cell = ws.cell(row=row, column=1, value=f"{p}. Ders")
            period_cell.fill = PERIOD_FILL
            period_cell.font = Font(bold=True, size=9)
            period_cell.alignment = Alignment(horizontal="center", vertical="center")
            period_cell.border = border

            for d in range(DAYS):
                day_max = inp.daily_periods[d]
                cell = ws.cell(row=row, column=d + 2)
                cell.border = border
                cell.alignment = Alignment(
                    horizontal="center", vertical="center", wrap_text=True
                )
                if p > day_max:
                    cell.value = "—"
                    cell.fill = PatternFill("solid", fgColor="E2E8F0")
                    cell.font = Font(color="94A3B8", size=9)
                elif (d, p) in grid:
                    cell.value = grid[(d, p)]
                    cell.fill = LESSON_FILL
                    cell.font = Font(size=9)
                else:
                    cell.value = ""
                    cell.fill = EMPTY_FILL

    # Sınıf sayfaları
    class_names = sorted(set(l.class_name for l in result.lessons))
    # Ayrıca inp.classes'dan isimler de ekle (dersi olmayan sınıflar için de sheet aç)
    all_class_names = sorted(set(cls.name for cls in inp.classes) | set(class_names))
    for cls_name in all_class_names:
        ws = wb.create_sheet(title=f"📚 {cls_name}")
        _write_timetable_sheet(ws, cls_name, "class")

    # Öğretmen sayfaları
    teacher_names = sorted(set(l.teacher_name for l in result.lessons))
    for t_name in teacher_names:
        safe = t_name[:28] if len(t_name) > 28 else t_name
        ws = wb.create_sheet(title=f"👨‍🏫 {safe}")
        _write_timetable_sheet(ws, t_name, "teacher")

    # Uyarılar sayfası
    if result.warnings:
        ws_warn = wb.create_sheet(title="⚠️ Uyarılar")
        ws_warn.column_dimensions["A"].width = 20
        ws_warn.column_dimensions["B"].width = 30
        ws_warn.column_dimensions["C"].width = 70

        headers = ["Sınıf", "Ders", "Açıklama"]
        for col, h in enumerate(headers, 1):
            cell = ws_warn.cell(row=1, column=col, value=h)
            cell.fill = PatternFill("solid", fgColor="DC2626")
            cell.font = Font(bold=True, color="FFFFFF", size=11)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border

        for row_idx, w in enumerate(result.warnings, 2):
            ws_warn.cell(row=row_idx, column=1, value=w.class_name).fill = WARN_FILL
            ws_warn.cell(row=row_idx, column=2, value=w.course_name).fill = WARN_FILL
            reason_cell = ws_warn.cell(row=row_idx, column=3, value=w.reason)
            reason_cell.fill = WARN_FILL
            reason_cell.alignment = Alignment(wrap_text=True)
            for col in range(1, 4):
                ws_warn.cell(row=row_idx, column=col).border = border

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def generate_pdf(inp: SchedulerInput, result: SchedulerResult) -> bytes:
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph,
        Spacer, PageBreak
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "title", parent=styles["Heading1"], alignment=TA_CENTER,
        textColor=colors.HexColor("#1E3A5F"), fontSize=14, spaceAfter=8
    )
    cell_style = ParagraphStyle(
        "cell", parent=styles["Normal"], fontSize=7, leading=9, alignment=TA_CENTER
    )

    story = []
    max_p = max(inp.daily_periods)

    def _make_table(entity_name: str, entity_type: str) -> Table:
        grid = _build_grid(entity_name, entity_type, result, inp.daily_periods)

        header_row = ["Ders"] + DAY_NAMES
        rows = [header_row]
        for p in range(1, max_p + 1):
            row = [Paragraph(f"<b>{p}. Ders</b>", cell_style)]
            for d in range(DAYS):
                day_max = inp.daily_periods[d]
                if p > day_max:
                    row.append(Paragraph("—", cell_style))
                elif (d, p) in grid:
                    row.append(Paragraph(grid[(d, p)].replace("\n", "<br/>"), cell_style))
                else:
                    row.append("")
            rows.append(row)

        col_widths = [2.5*cm] + [4.8*cm] * DAYS
        tbl = Table(rows, colWidths=col_widths, rowHeights=[0.7*cm] + [1.5*cm] * max_p)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, 0), 9),
            ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#EFF6FF")),
            ("FONTNAME",   (0, 1), (0, -1), "Helvetica-Bold"),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ]))
        return tbl

    all_class_names = sorted(set(cls.name for cls in inp.classes))
    for cls_name in all_class_names:
        story.append(Paragraph(f"Sınıf Programı: {cls_name}", title_style))
        story.append(_make_table(cls_name, "class"))
        story.append(PageBreak())

    teacher_names = sorted(set(l.teacher_name for l in result.lessons))
    for t_name in teacher_names:
        story.append(Paragraph(f"Öğretmen Programı: {t_name}", title_style))
        story.append(_make_table(t_name, "teacher"))
        story.append(PageBreak())

    # Uyarılar
    if result.warnings:
        story.append(Paragraph("⚠️ Uyarılar ve Atanamayan Dersler", title_style))
        warn_data = [["Sınıf", "Ders", "Açıklama"]]
        for w in result.warnings:
            warn_data.append([w.class_name, w.course_name, w.reason])
        warn_tbl = Table(warn_data, colWidths=[3*cm, 5*cm, 17*cm])
        warn_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DC2626")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#FEF3C7")),
            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN",      (0, 0), (1, -1), "CENTER"),
        ]))
        story.append(warn_tbl)

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# WORD
# ---------------------------------------------------------------------------

def generate_word(inp: SchedulerInput, result: SchedulerResult) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _cell_para(cell, text: str, bold=False, color="000000", size=9, center=False):
        cell.text = ""
        para = cell.paragraphs[0]
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)

    doc = Document()
    # Sayfa yatay
    from docx.oxml import OxmlElement as OE
    section = doc.sections[0]
    section.orientation = 1  # landscape
    section.page_width, section.page_height = section.page_height, section.page_width

    max_p = max(inp.daily_periods)

    def _write_entity(entity_name: str, entity_type: str, label: str):
        grid = _build_grid(entity_name, entity_type, result, inp.daily_periods)
        heading = doc.add_heading(f"{label}: {entity_name}", level=2)
        heading.runs[0].font.color.rgb = RGBColor.from_string("1E3A5F")

        table = doc.add_table(rows=max_p + 1, cols=DAYS + 1)
        table.style = "Table Grid"

        # Başlık satırı
        header_cells = table.rows[0].cells
        _cell_para(header_cells[0], "Ders / Gün", bold=True, color="FFFFFF", center=True)
        _set_cell_bg(header_cells[0], "1E3A5F")
        for d, day_name in enumerate(DAY_NAMES):
            _cell_para(header_cells[d + 1], day_name, bold=True, color="FFFFFF", center=True)
            _set_cell_bg(header_cells[d + 1], "2563EB")

        for p in range(1, max_p + 1):
            row = table.rows[p]
            _cell_para(row.cells[0], f"{p}. Ders", bold=True, center=True, size=8)
            _set_cell_bg(row.cells[0], "EFF6FF")

            for d in range(DAYS):
                day_max = inp.daily_periods[d]
                cell = row.cells[d + 1]
                if p > day_max:
                    _cell_para(cell, "—", color="94A3B8", center=True, size=8)
                    _set_cell_bg(cell, "E2E8F0")
                elif (d, p) in grid:
                    text = grid[(d, p)]
                    _cell_para(cell, text, center=True, size=8)
                    _set_cell_bg(cell, "DBEAFE")
                else:
                    _set_cell_bg(cell, "F8FAFC")

        doc.add_paragraph()
        doc.add_page_break()

    all_class_names = sorted(set(cls.name for cls in inp.classes))
    for cls_name in all_class_names:
        _write_entity(cls_name, "class", "Sınıf Programı")

    teacher_names = sorted(set(l.teacher_name for l in result.lessons))
    for t_name in teacher_names:
        _write_entity(t_name, "teacher", "Öğretmen Programı")

    # Uyarılar
    if result.warnings:
        doc.add_heading("⚠️ Uyarılar ve Atanamayan Dersler", level=2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdrs = ["Sınıf", "Ders", "Açıklama"]
        for i, h in enumerate(hdrs):
            _cell_para(tbl.rows[0].cells[i], h, bold=True, color="FFFFFF", center=True)
            _set_cell_bg(tbl.rows[0].cells[i], "DC2626")
        for w in result.warnings:
            row = tbl.add_row()
            _cell_para(row.cells[0], w.class_name, center=True, size=8)
            _cell_para(row.cells[1], w.course_name, center=True, size=8)
            _cell_para(row.cells[2], w.reason, size=8)
            for c in row.cells:
                _set_cell_bg(c, "FEF3C7")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Ana dispatch fonksiyonu
# ---------------------------------------------------------------------------

def generate_output(
    inp: SchedulerInput,
    result: SchedulerResult,
    fmt: OutputFormat,
) -> tuple[bytes, str, str]:
    """
    Döner: (bytes, media_type, filename)
    """
    if fmt == "excel":
        data = generate_excel(inp, result)
        return data, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "ders_programi.xlsx"
    elif fmt == "pdf":
        data = generate_pdf(inp, result)
        return data, "application/pdf", "ders_programi.pdf"
    elif fmt == "word":
        data = generate_word(inp, result)
        return data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "ders_programi.docx"
    else:
        raise ValueError(f"Bilinmeyen format: {fmt}")
